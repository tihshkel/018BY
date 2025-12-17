#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Скрипт для создания Word документа из отчета о нововведениях
с оформлением по ГОСТу
"""

import os
import sys
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.section import WD_SECTION
except ImportError:
    print("Ошибка: python-docx не установлен.")
    print("Установите: pip install python-docx")
    sys.exit(1)


def setup_page_margins(section):
    """Настройка полей страницы по ГОСТу"""
    # ГОСТ: верхнее и нижнее - 2 см, левое - 3 см, правое - 1.5 см
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)


def setup_paragraph_style(paragraph, font_size=14, bold=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=True):
    """Настройка стиля абзаца по ГОСТу"""
    paragraph.alignment = alignment
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph_format.space_after = Pt(0)
    if first_line_indent:
        paragraph_format.first_line_indent = Cm(1.25)  # Красная строка 1.25 см
    else:
        paragraph_format.first_line_indent = Cm(0)
    
    # Если в параграфе уже есть runs, используем первый, иначе создаем новый
    if paragraph.runs:
        run = paragraph.runs[0]
    else:
        run = paragraph.add_run('')
    
    run.font.name = 'Times New Roman'
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def create_title_page(doc):
    """Создание титульного листа"""
    # Добавляем титульный лист
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Название организации (если нужно)
    org_run = title_paragraph.add_run('ОТЧЕТ О НОВОВВЕДЕНИЯХ')
    org_run.font.name = 'Times New Roman'
    org_run.font.size = Pt(16)
    org_run.font.bold = True
    
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run('В ПРОЕКТЕ 018BY')
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    
    # Период
    doc.add_paragraph()  # Пустая строка
    period_paragraph = doc.add_paragraph()
    period_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    period_run = period_paragraph.add_run('Период: последние две недели')
    period_run.font.name = 'Times New Roman'
    period_run.font.size = Pt(14)
    
    # Дата
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_paragraph.add_run(f'Дата составления: {datetime.now().strftime("%d.%m.%Y")}')
    date_run.font.name = 'Times New Roman'
    date_run.font.size = Pt(14)
    
    # Разрыв страницы
    doc.add_page_break()


def create_word_document(input_file, output_file):
    """Создание Word документа из текстового файла"""
    
    # Читаем исходный файл
    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Создаем документ
    doc = Document()
    
    # Настраиваем поля первой секции
    section = doc.sections[0]
    setup_page_margins(section)
    
    # Создаем титульный лист
    create_title_page(doc)
    
    # Настраиваем поля для остальных страниц
    for section in doc.sections[1:]:
        setup_page_margins(section)
    
    # Разбиваем содержимое на абзацы
    paragraphs = content.split('\n\n')
    
    current_section = None
    is_first_paragraph = True
    
    for para_text in paragraphs:
        para_text = para_text.strip()
        
        if not para_text:
            continue
        
        # Определяем тип абзаца
        if para_text.startswith('ОТЧЕТ') or para_text.startswith('Период') or para_text.startswith('Дата'):
            # Пропускаем заголовок, так как он уже на титульном листе
            continue
        elif para_text.startswith('ОБЩЕЕ ОПИСАНИЕ') or para_text.startswith('ОСНОВНЫЕ НОВОВВЕДЕНИЯ'):
            # Заголовок раздела
            current_section = para_text
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(para_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(12)
            paragraph_format.space_after = Pt(12)
            paragraph_format.first_line_indent = Cm(0)  # Без красной строки для заголовков
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            doc.add_paragraph()  # Пустая строка после заголовка
            is_first_paragraph = True
        else:
            # Обычный абзац
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(para_text)
            
            # Для первого абзаца после заголовка убираем красную строку
            if is_first_paragraph and current_section:
                setup_paragraph_style(paragraph, font_size=14, bold=False, first_line_indent=False)
                is_first_paragraph = False
            else:
                setup_paragraph_style(paragraph, font_size=14, bold=False, first_line_indent=True)
    
    # Добавляем нумерацию страниц (начиная со второй страницы)
    # Примечание: python-docx не поддерживает напрямую нумерацию,
    # но можно добавить колонтитулы
    
    # Сохраняем документ
    doc.save(output_file)
    print(f"✓ Документ Word создан: {output_file}")


def main():
    # Определяем пути
    script_dir = Path(__file__).parent
    project_root = script_dir.parent
    
    input_file = project_root / 'docs' / 'отчет_нововведения_2_недели.txt'
    output_file = project_root / 'docs' / 'отчет_нововведения_2_недели.docx'
    
    if not input_file.exists():
        print(f"Ошибка: файл {input_file} не найден")
        sys.exit(1)
    
    print("Создание Word документа из отчета...")
    print(f"Входной файл: {input_file}")
    print(f"Выходной файл: {output_file}")
    
    try:
        create_word_document(input_file, output_file)
        print(f"✓ Готово! Документ сохранен: {output_file}")
    except Exception as e:
        print(f"Ошибка при создании документа: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()

